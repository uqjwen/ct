#!/usr/bin/env python3
"""Generate the line annotation, analysis model, and verification specification."""

from __future__ import annotations

import hashlib
import html
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ct_lsu_ld_ag.v"
ANNOTATION = ROOT / "ct_lsu_ld_ag.md"
SPEC = ROOT / "ld_ag_spec.docx"
MODEL_DIR = ROOT / "work" / "analysis"
MODEL = MODEL_DIR / "rtl_model.json"
UPSTREAM_URL = (
    "https://github.com/XUANTIE-RV/openc910/blob/main/"
    "C910_RTL_FACTORY/gen_rtl/lsu/rtl/ct_lsu_ld_ag.v"
)


SECTIONS = [
    (1, 15, "许可证与版权"),
    (16, 160, "模块端口列表"),
    (161, 303, "端口方向与位宽声明"),
    (304, 541, "内部寄存器、连线与参数"),
    (542, 569, "RF 映射与门控时钟"),
    (570, 778, "AG 流水寄存器与重放数据保持"),
    (779, 811, "虚拟地址生成"),
    (812, 835, "指令类型与预取分类"),
    (836, 910, "访问大小、对齐与 16-byte 边界"),
    (911, 1030, "字节有效掩码与旋转"),
    (1031, 1060, "MMU 接口"),
    (1061, 1105, "物理地址、边界加速与转发"),
    (1106, 1120, "提交匹配"),
    (1121, 1197, "DCACHE tag/data 阵列请求"),
    (1198, 1242, "异常生成"),
    (1243, 1337, "stall、restart 与优先级"),
    (1338, 1397, "下传 DC 与提前唤醒"),
    (1398, 1443, "本地监控、LSIQ 与 IDU 接口"),
    (1444, 1462, "性能计数与模块结束"),
]


SIGNAL_DESCRIPTIONS = {
    "cpurst_b": "低有效异步复位。",
    "forever_cpuclk": "未门控的 CPU 基准时钟，作为 AG ICG 输入。",
    "ctrl_ld_clk": "load 控制时钟，用于更新 AG valid。",
    "cp0_yy_clk_en": "全局时钟使能。",
    "cp0_lsu_icg_en": "LSU 模块时钟门控使能。",
    "pad_yy_icg_scan_en": "扫描模式下的 ICG 覆盖使能。",
    "cp0_lsu_dcache_en": "数据缓存使能；门控 tag/data array 请求。",
    "cp0_lsu_cb_aclr_dis": "禁止跨 16-byte 边界的 cache-buffer acceleration。",
    "cp0_lsu_da_fwd_dis": "禁止 load ahead/DA forwarding 相关有效信号。",
    "cp0_lsu_mm": "misaligned load 处理模式配置；为 0 时普通 load 不对齐在 AG 报错。",
    "idu_lsu_rf_pipe3_sel": "RF pipe3 向 AG 发射/选择有效指示。",
    "idu_lsu_rf_pipe3_gateclk_sel": "RF pipe3 门控时钟选择；本模块同时把它视为 RF 数据有效。",
    "idu_lsu_rf_pipe3_src0": "地址基址源操作数。",
    "idu_lsu_rf_pipe3_src1": "LDR 类指令的 offset 源操作数。",
    "idu_lsu_rf_pipe3_offset": "普通 load 的 12-bit 立即数 offset。",
    "idu_lsu_rf_pipe3_offset_plus": "split 首拍使用的 13-bit 备选 offset。",
    "idu_lsu_rf_pipe3_shift": "offset 移位选择，RTL 按 one-hot 编码使用。",
    "idu_lsu_rf_pipe3_inst_size": "访问大小编码：BYTE/HALF/WORD/DWORD。",
    "idu_lsu_rf_pipe3_inst_type": "load/atomic 子类型编码。",
    "idu_lsu_rf_pipe3_atomic": "原子类指令标志。",
    "idu_lsu_rf_pipe3_inst_ldr": "寄存器 offset 的 LDR 类地址生成标志。",
    "idu_lsu_rf_pipe3_unalign_2nd": "split/边界访问第二拍标志，锁存为 ld_ag_secd。",
    "idu_lsu_rf_pipe3_lch_entry": "对应 LSIQ entry 的 12-bit one-hot/bitmap。",
    "idu_lsu_rf_pipe3_iid": "指令 ID，用于提交和新旧比较。",
    "mmu_lsu_pa0": "MMU 返回的物理页号。",
    "mmu_lsu_pa0_vld": "MMU port0 返回有效；同时限定 CA/SO 和多种异常。",
    "mmu_lsu_page_fault0": "MMU port0 页故障。",
    "mmu_lsu_stall0": "MMU stall；上游当前 DUTLB 实现中常为 0，但本模块未按 valid 门控。",
    "mmu_lsu_ca0": "页属性：cacheable。",
    "mmu_lsu_so0": "页属性：strong order。",
    "dcache_arb_ag_ld_sel": "DCACHE arbiter 已选择 load AG 请求；取反形成 cache stall。",
    "dcache_arb_ld_ag_borrow_addr_vld": "DCACHE arbiter 借用地址有效，覆盖下传 DC 地址。",
    "dcache_arb_ld_ag_addr": "DCACHE arbiter 借用的地址。",
    "rtu_yy_xx_flush": "全局 flush，清除 AG valid 并中止 MMU。",
    "ld_ag_inst_vld": "AG 当前指令有效；stall 时保持，flush/reset 时清零。",
    "ld_ag_dc_inst_vld": "发往 DC 的有效，排除所有 stall/restart。",
    "lsu_mmu_va0_vld": "发往 MMU port0 的请求有效，等于 AG valid。",
    "lsu_mmu_va0": "发往 MMU 的 base VA；页内 offset 在 AG 本地加入。",
    "lsu_mmu_abort0": "中止当前 MMU port0 请求的组合条件。",
    "ld_ag_pa": "MMU PPN 与最终 VA 低 12 位拼接的物理地址。",
    "ld_ag_boundary": "普通 load 跨 16-byte 边界或处于第二拍。",
    "ld_ag_dc_bytes_vld": "当前 16-byte 数据窗口的字节有效掩码。",
    "ld_ag_dc_bytes_vld1": "边界访问低地址侧/合并用途的辅助掩码。",
    "ld_ag_dc_rot_sel": "基于 VA 低 4 位的数据旋转选择。",
    "ld_ag_dc_acclr_en": "满足 PA valid、cacheable 等条件的边界加速使能。",
    "ld_ag_expt_vld": "聚合异常：misalign、SO misalign、带页属性 access fault、page fault。",
    "ld_ag_expt_ldamo_not_ca": "原子 load/AMO 访问 non-cacheable 页的专用异常。",
    "ld_ag_stall_ori": "不含 atomic-no-commit restart 的原始 stall 指示。",
    "ld_ag_stall_restart_entry": "stall/restart 时选中的 LSIQ entry bitmap。",
    "ld_ag_ahead_predict": "load ahead 预测，本版本恒为 1。",
}


IMPORTANT_LINES = {
    551: "生成 AG 数据门控时钟的 local enable：新 RF 指令或 AG 已有有效指令时开钟。",
    582: "AG valid 的时序逻辑，使用独立的 ctrl_ld_clk 和低有效异步复位。",
    588: "若 AG stall 或 RF 有新指令，则下一周期 AG valid 保持/置 1；该条件依赖 stall 不在空闲期误拉高。",
    626: "AG 数据寄存器组的时序逻辑，使用门控后的 ld_ag_clk。",
    660: "仅在 AG 不 stall 且 RF 数据有效时装载新指令数据；stall 时自然保持。",
    698: "记录一次跨页/边界 LDR 重放已经发生，避免重复执行同一重放动作。",
    713: "保存 offset 的 one-hot 移位选择；跨页 LDR 重放时重置为不移位。",
    728: "更新 offset 高 32 位；无 reset，必须与有效位一起解释。",
    741: "更新 offset 低 32 位；边界第二拍首次重放时注入 0x10。",
    757: "保存 split 首拍的 offset_plus，跨页重放时清零。",
    771: "保存 base；跨页重放时把已算出的 VA 回写为新 base。",
    789: "按 one-hot shift[3:0] 对 offset 选择左移 0/1/2/3 位并 OR 合成。",
    794: "计算常规虚拟地址 va_ori = base + shifted_offset。",
    797: "计算 split 首拍候选虚拟地址 va_plus = base + sign_extended(offset_plus)。",
    801: "普通非 LDR load 首拍跨 16-byte 边界时选择 va_plus。",
    806: "在 va_plus 与 va_ori 之间选择本拍最终虚拟地址。",
    816: "普通 load 分类：非 atomic 且 inst_type 为 00。",
    819: "LR 分类：atomic 且 inst_type 为 01。",
    821: "LDAMO 分类：atomic 且 inst_type 为 00。",
    842: "把 2-bit 指令大小译码为 size-minus-one 掩码 0/1/3/7。",
    872: "根据 size 和 VA 低 3 位判断自然对齐。",
    903: "VA 低 4 位加 size-minus-one；第 4 位进位表示跨 16-byte 边界。",
    907: "边界输出仅对普通 load 有效；第二拍强制继续标记 boundary。",
    915: "生成从起始 byte 到 16-byte 窗口末端的高侧全掩码。",
    940: "生成从窗口起点到访问结束 byte 的低侧全掩码。",
    964: "非跨界访问的精确 byte mask 等于高侧掩码与低侧掩码的交集。",
    983: "第二拍选择高地址侧掩码，首拍/普通拍选择低侧或交集掩码。",
    1035: "MMU 请求有效等于 AG valid。",
    1036: "MMU 请求的是 base 地址而不是最终 VA；跨页由重放机制处理。",
    1038: "组合生成 MMU abort：跨页重放、atomic 等提交、DCACHE stall、misalign 或 flush。",
    1049: "SO 属性只在 MMU PA valid 时对外有效。",
    1051: "cacheable 属性只在 MMU PA valid 时对外有效。",
    1056: "uTLB miss 直接定义为 PA 无效；空闲态也可能为 1，需与指令 valid 合用。",
    1065: "物理地址由 MMU 返回 PPN 和最终 VA 的 12-bit 页内 offset 拼接。",
    1072: "形成边界加速预资格，不包含 MMU PA valid/cacheable；真实下传使能在后续补充。",
    1081: "真实 DC 边界加速使能还要求 PA valid 且页面 cacheable。",
    1102: "forward bypass 只允许受支持大小、非 vector-memory、非 boundary 访问。",
    1109: "比较 RTU commit0 的 valid+iid 与当前 AG iid。",
    1128: "DCACHE 阵列请求只由 AG valid 和 dcache enable 门控，属于投机读取。",
    1138: "tag index 使用 PA[14:6]。",
    1145: "根据 boundary/secd 和 4-byte bank 起止位置生成低半区 bank enable。",
    1175: "边界加速时强制访问全部四个 32-bit bank。",
    1189: "低 128-bit 数据阵列索引使用 PA[14:4]。",
    1190: "高 128-bit 数据阵列索引翻转 PA[4]，指向相邻 16-byte 块。",
    1204: "生成无需等待页属性的 misalign：atomic/vector 或普通 load 且 MM 模式关闭。",
    1211: "SO 页面上的普通 load 不对齐异常，需要 PA valid。",
    1217: "页故障输出要求 PA valid 与 MMU page_fault 同时成立。",
    1221: "生成 LDAMO 访问 non-cacheable 页的专用异常。",
    1238: "聚合异常不包含专用 LDAMO-not-cacheable 异常，消费者需另行处理。",
    1247: "atomic 指令在对应 iid 尚未 commit 时请求 restart。",
    1256: "用 base 页内低 12 位与有效 offset 低位计算 4 KiB 进位。",
    1262: "检测 shifted offset 高位是否全 0。",
    1263: "检测 shifted offset 高位是否全 1，用于合法负 offset 的符号扩展判定。",
    1270: "跨 4 KiB 条件为低位加法进位或 offset 高位既非全 0 也非全 1。",
    1279: "跨页或 LDR 边界第二拍需要 stall/replay，并排除已在 AG 报出的 misalign。",
    1283: "DCACHE 未选择当前 AG 请求且指令有效时产生 cache stall。",
    1285: "MMU stall 未在本模块按 AG valid 门控；当前上游 DUTLB 实现将其常置 0。",
    1298: "总 restart/stall 合并 atomic、跨页、DCACHE 和 MMU 四类原因。",
    1304: "跨页/LDR replay 仅在非 atomic 优先级且未被更老 RF 指令 mask 时执行。",
    1310: "对外原始 stall 排除 atomic-no-commit restart。",
    1321: "实例化 iid 比较器，判断 RF 指令是否比当前 AG 指令更老。",
    1331: "更老 RF 指令可 mask 当前 AG stall，以改善调度性能。",
    1342: "下传 DC valid = AG valid 且不存在任何 stall/restart。",
    1345: "DC 伴随 mmu_req 为 abort 的反相；空闲时可能为 1，DC 仅在 AG valid 时锁存。",
    1348: "DCACHE borrow 地址有效时覆盖正常的 AG 物理地址。",
    1357: "ahead prediction 在本版本恒为 1，且源码留有 TODO。",
    1364: "普通 load ahead-valid 排除 boundary、vector 和 DA-forward-disable。",
    1371: "vector load valid 使用 inst_vfls 分类，且排除 VMB merge 与不能加速的 boundary。",
    1382: "vector load ahead-valid 在本版本恒为 0。",
    1388: "比较 store AG iid 与 load AG iid，生成 RAW 新旧关系信号。",
    1405: "atomic 指令已经 commit 时初始化 local monitor。",
    1416: "atomic 未 commit restart 时向 IDU/LSIQ 提供 wait-old 门控。",
    1423: "为 IDU timing 生成普通 load 的提前有效。",
    1434: "为 IDU timing 生成 vector load 的提前有效。",
    1447: "生成跨 4 KiB stall 的性能事件脉冲。",
    1452: "生成其他 stall 的性能事件，排除跨 4 KiB 与 uTLB miss。",
}


def section_for(line_no: int) -> str:
    for start, end, name in SECTIONS:
        if start <= line_no <= end:
            return name
    return "其他"


def signal_description(name: str) -> str:
    if name in SIGNAL_DESCRIPTIONS:
        return SIGNAL_DESCRIPTIONS[name]
    prefix_rules = [
        ("idu_lsu_rf_pipe3_", "来自 IDU/RF pipe3 的指令属性或操作数。"),
        ("mmu_lsu_", "来自 MMU port0 的翻译结果、属性或流控。"),
        ("lsu_mmu_", "发往 MMU port0 的请求、标识或中止信息。"),
        ("ag_dcache_arb_", "发往 DCACHE arbiter 的阵列请求、索引或门控。"),
        ("dcache_arb_", "来自 DCACHE arbiter 的选择或借用地址信息。"),
        ("ld_ag_dc_", "从 AG 下传到 load DC stage 的数据或控制。"),
        ("ld_ag_expt_", "AG 生成的异常条件。"),
        ("ld_ag_page_", "MMU 页面属性的 AG 侧输出。"),
        ("ld_ag_", "load AG stage 的状态、地址、分类或控制输出。"),
        ("lsu_idu_", "从 LSU/AG 返回 IDU 的调度或目的寄存器信息。"),
        ("rtu_yy_xx_commit", "来自 RTU 的提交端口信息。"),
        ("cp0_", "来自 CP0 的 LSU 配置或时钟控制。"),
    ]
    for prefix, desc in prefix_rules:
        if name.startswith(prefix):
            return desc
    return "模块接口信号；具体有效条件由相邻控制信号和功能规则限定。"


def parse_ports(lines: list[str]) -> list[dict[str, str]]:
    ports: list[dict[str, str]] = []
    pattern = re.compile(r"^(input|output)\s+(?:\[([^\]]+)\]\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*;")
    for idx, raw in enumerate(lines, 1):
        match = pattern.match(raw.strip())
        if not match:
            continue
        direction, width, name = match.groups()
        ports.append(
            {
                "line": str(idx),
                "direction": direction,
                "width": width.replace(" ", "") if width else "1",
                "name": name,
                "description": signal_description(name),
            }
        )
    return ports


def explain_line(line_no: int, raw: str, current_section: str) -> str:
    stripped = raw.strip()
    if line_no in IMPORTANT_LINES:
        return IMPORTANT_LINES[line_no]
    if not stripped:
        return f"空行，用于分隔“{current_section}”中的逻辑块，提高源码可读性。"
    if stripped.startswith("/*"):
        return "Apache-2.0 版权/许可证块开始。"
    if stripped == "*/":
        return "版权/许可证块结束。"
    if line_no <= 14:
        return "Apache-2.0 许可证原文，规定本 RTL 的使用、分发和免责声明。"
    if stripped.startswith("//"):
        text = stripped[2:].strip()
        if "&" in text and "@" in text:
            return "源代码生成工具保留的指令/行号标记；不参与综合逻辑。"
        if not text or set(text) <= {"=", "-", "+"}:
            return f"“{current_section}”的视觉分隔注释，不参与逻辑。"
        return f"原作者注释，说明“{text}”；仅用于解释相邻 RTL，不参与综合。"
    if stripped.startswith("module "):
        return "声明顶层模块 ct_lsu_ld_ag，并开始列举外部端口。"
    if stripped == ");" and line_no < 200:
        return "结束模块端口列表。"
    if stripped == "endmodule":
        return "结束 ct_lsu_ld_ag 模块定义。"

    decl = re.match(
        r"^(input|output|wire|reg)\s+(?:\[([^\]]+)\]\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*;",
        stripped,
    )
    if decl:
        kind, width, name = decl.groups()
        kind_cn = {"input": "输入", "output": "输出", "wire": "组合连线", "reg": "过程赋值变量"}[kind]
        width_cn = f"位段 [{width.replace(' ', '')}]" if width else "1 bit"
        return f"声明{kind_cn} `{name}`，{width_cn}。{signal_description(name)}"
    if stripped.startswith("parameter ") or (533 <= line_no <= 540 and "=" in stripped):
        return "声明/延续本模块的本地参数编码，用于访问大小、队列深度或 PC 位宽。"
    if 17 < line_no < 159 and re.match(r"[A-Za-z_][A-Za-z0-9_]*,?$", stripped):
        return f"将 `{stripped.rstrip(',')}` 列入模块端口列表。{signal_description(stripped.rstrip(','))}"

    assign = re.match(r"^assign\s+([^=]+?)\s*=", stripped)
    if assign:
        lhs = assign.group(1).strip()
        return f"连续组合赋值开始：驱动 `{lhs}`；完整右值可能延续到后续行，属于“{current_section}”。"
    if stripped.startswith("always @"):
        if "posedge" in stripped or "negedge" in stripped:
            return "开始时序 always 块；敏感表给出采样时钟及可能的异步复位。"
        return "开始组合 always 块；敏感表列出决定本块输出的输入信号。"
    if stripped.startswith("casez"):
        return "开始带通配位的 casez 译码；`?` 位不参与匹配。"
    if stripped.startswith("case"):
        return "开始 case 译码，根据括号内编码选择唯一分支。"
    if stripped == "endcase":
        return "结束当前 case/casez 译码。"
    if stripped == "begin":
        return "开始上一条 always/if/case 语句的复合语句块。"
    if stripped == "end":
        return "结束当前复合语句块。"
    if stripped.startswith("if") or stripped.startswith("else if"):
        return f"条件分支：仅当 `{stripped}` 的条件成立时执行紧随的赋值/语句。"
    if stripped == "else":
        return "前述条件均不成立时进入默认分支。"
    if "<=" in stripped:
        lhs = stripped.split("<=", 1)[0].strip()
        return f"时序非阻塞赋值：在时钟沿后更新 `{lhs}`；右值及条件由本行和相邻分支定义。"
    if re.match(r"^(default:|\{?[0-9A-Z_].*:)", stripped) and "=" in stripped:
        return "case 译码分支：为当前输入组合设置确定的组合输出。"
    if stripped.startswith("gated_clk_cell"):
        return "实例化通用门控时钟单元，生成 load AG 数据路径时钟。"
    if stripped.startswith("ct_rtu_compare_iid"):
        return "实例化 RTU iid 新旧比较器。"
    inst_port = re.match(r"^\.([A-Za-z_][A-Za-z0-9_]*)\s*\(([^)]*)\)", stripped.rstrip(","))
    if inst_port:
        port, sig = inst_port.groups()
        return f"实例端口连接：子模块端口 `{port}` 连接本层信号/常量 `{sig.strip()}`。"
    if stripped == ");":
        return "结束当前子模块实例的端口连接。"
    if stripped.startswith("|") or stripped.startswith("+"):
        return "注释中的 ASCII 框图行，用来列举流水寄存器字段，不参与逻辑。"
    if stripped.endswith(";") or stripped.endswith(",") or stripped.startswith(("&&", "||", "?", ":", "+", "==", "&", "|", ".")):
        return f"延续“{current_section}”中的声明、实例连接或多行组合表达式；需与前后行合并阅读。"
    return f"“{current_section}”中的 RTL 语句片段；与相邻行共同构成完整声明、条件或表达式。"


def build_annotation(lines: list[str], digest: str) -> None:
    out = [
        "# `ct_lsu_ld_ag.v` 逐行注释",
        "",
        "## 文档目的",
        "",
        "本文对 `ct_lsu_ld_ag.v` 的 1462 个物理行逐行建立锚点，并给出中文解释。解释严格区分代码事实和接口推断；涉及时序契约的疑点请同时阅读 `ld_ag_spec.docx` 与 `ct_lsu_ld_ag_risk_review.md`。",
        "",
        "## 源码基线",
        "",
        f"- SHA-256：`{digest}`",
        f"- 上游同源文件：<{UPSTREAM_URL}>",
        "- 覆盖规则：每个物理行（包括空行）在下表中出现一次，行号与源码一一对应。",
        "",
        "## 功能总览",
        "",
        "`ct_lsu_ld_ag` 是 LSU load 流水线的 AG（address generation）级。它接收 RF pipe3 的基址、offset、大小和指令属性，生成 VA/PA、字节掩码、DCACHE 阵列索引、MMU 请求、异常、stall/restart、load-ahead 以及下游 DC/IDU/LSIQ 伴随信号。",
        "",
        "核心流程：`RF 发射 -> AG 寄存 -> VA/边界判断 -> MMU 页翻译 -> PA/阵列索引 -> 异常与 stall 仲裁 -> DC 下传`。",
        "",
        "## 分段索引",
        "",
        "| 源码范围 | 功能块 |",
        "|---|---|",
    ]
    for start, end, name in SECTIONS:
        out.append(f"| L{start:04d}-L{end:04d} | {name} |")

    current = None
    for idx, raw in enumerate(lines, 1):
        section = section_for(idx)
        if section != current:
            current = section
            out.extend(
                [
                    "",
                    f"## {section}",
                    "",
                    "| 行号 | 原始代码 | 中文注释 |",
                    "|---:|---|---|",
                ]
            )
        source_cell = "&nbsp;" if not raw else f"<code>{html.escape(raw).replace('|', '&#124;')}</code>"
        note = explain_line(idx, raw, section).replace("|", "&#124;")
        out.append(f"| <!-- SRC-LINE:{idx:04d} -->L{idx:04d} | {source_cell} | {note} |")
    out.extend(
        [
            "",
            "## 阅读注意事项",
            "",
            "1. `ld_ag_inst_vld` 是大多数数据/属性的总有效资格；若某个输出自身未门控，消费者仍应按对应 valid 采样。",
            "2. MMU 接收 base VA，AG 在本地加入页内 offset；跨 4 KiB 必须依赖 replay 重新翻译。",
            "3. DCACHE tag/data 阵列请求是投机的，不能等价为 load 已经可以提交。",
            "4. `ld_ag_ahead_predict` 在当前版本恒为 1，后级必须具备完整取消/回放协议。",
        ]
    )
    ANNOTATION.write_text("\n".join(out) + "\n", encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], "17365D")
        set_cell_text(hdr.cells[i], header, bold=True, color="FFFFFF", size=8.5)
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2:
            for cell in cells:
                set_cell_shading(cell, "EAF0F6")
        for i, value in enumerate(values):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = document.add_paragraph(style=style)
    p.add_run(text)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.add_run(text)


def add_code(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Code Block")
    p.add_run(text)


def build_spec(lines: list[str], ports: list[dict[str, str]], digest: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in [
        ("Title", 28, "17365D"),
        ("Heading 1", 18, "17365D"),
        ("Heading 2", 14, "2F5597"),
        ("Heading 3", 11.5, "365F91"),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Cm(0.7)
    code_style.paragraph_format.right_indent = Cm(0.7)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("CT LSU Load Address Generation Verification Specification")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated from reviewed RTL baseline - 2026-07-11")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(110, 110, 110)

    # Cover page
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CT LSU Load Address Generation")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string("17365D")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Verification Specification")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("2F5597")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.add_run("Module: ct_lsu_ld_ag\n").bold = True
    p.add_run("Version: 1.0\n")
    p.add_run("Status: RTL-derived draft for verification review\n")
    p.add_run("Date: 2026-07-11")
    doc.add_paragraph()
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["RTL file", "ct_lsu_ld_ag.v"],
            ["RTL SHA-256", digest],
            ["Physical lines", str(len(lines))],
            ["Upstream baseline", UPSTREAM_URL],
            ["Important limitation", "Standalone module only; adjacent-stage timing is inferred from signal usage and must be confirmed by design."],
        ],
        [4.2, 12.0],
    )
    doc.add_page_break()

    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Revision", "Date", "Author", "Description"],
        [["1.0", "2026-07-11", "Codex", "Initial RTL-derived verification specification"]],
        [2.0, 2.8, 3.0, 8.4],
    )
    doc.add_heading("Table of Contents", level=1)
    for item in [
        "1. Purpose and Scope",
        "2. Module Overview",
        "3. Clock, Reset, and Pipeline Timing",
        "4. Interface Specification",
        "5. Functional Requirements",
        "6. Configuration Matrix",
        "7. Verification Strategy",
        "8. Suggested Assertions",
        "9. Coverage Model",
        "10. Open Questions and Assumptions",
    ]:
        add_number(doc, item)
    doc.add_page_break()

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This document specifies observable behavior of the load address-generation (LD AG) stage for block-level and subsystem verification. It is derived from the supplied RTL, not from a separate architecture document. Therefore, each statement is classified as an RTL fact, an inferred interface contract, or an open question."
    )
    doc.add_heading("1.1 In Scope", level=2)
    for text in [
        "RF-to-AG instruction capture and stall retention.",
        "Virtual-address calculation for immediate and register-offset loads.",
        "Natural-alignment, 16-byte boundary, and 4 KiB page-crossing detection.",
        "Byte masks, data rotation, DCACHE bank enables, and tag/data indices.",
        "MMU port0 request/abort behavior and PA construction.",
        "Exception, atomic commit, stall/restart, LSIQ, IDU timing, and performance outputs.",
        "Reset, clock-gating, invalid-state, and X-propagation expectations.",
    ]:
        add_bullet(doc, text)
    doc.add_heading("1.2 Out of Scope", level=2)
    for text in [
        "Correctness of the generic gated_clk_cell and ct_rtu_compare_iid implementations.",
        "DC, DA, WB, MMU, DCACHE, RTU, and IDU behavior except where their interface contract is required to verify AG.",
        "Vector mask generation removed/tied off in this RTL version.",
        "Performance targets, physical timing closure, and power sign-off.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("2. Module Overview", level=1)
    doc.add_paragraph(
        "The module accepts a load-like instruction from IDU/RF pipe3. Control valid is updated on ctrl_ld_clk, while instruction data is captured on a locally gated clock derived from forever_cpuclk. Address generation calculates an effective VA, detects 16-byte and 4 KiB crossings, requests page translation for the base page, constructs a PA, drives speculative DCACHE array reads, and decides whether the instruction advances to DC or restarts/stalls."
    )
    add_code(doc, "RF pipe3 -> AG registers -> VA/size/alignment -> MMU page result -> PA/masks/DCACHE index -> exception & stall arbiter -> DC/IDU/LSIQ")
    add_table(
        doc,
        ["Block", "Primary RTL lines", "Responsibility"],
        [
            ["Capture and hold", "L545-L777", "Capture valid/data, retain on stall, update replay base/offset."],
            ["VA generation", "L789-L810", "Shift offset, select normal/split VA, expose VPN."],
            ["Size/alignment/boundary", "L842-L909", "Decode access size and detect misalignment/16-byte crossing."],
            ["Masks and banks", "L915-L1030, L1145-L1190", "Generate byte masks, rotation, bank enables, and array indices."],
            ["MMU and PA", "L1035-L1068", "Issue base-page translation, consume attributes, build final PA."],
            ["Exceptions", "L1204-L1241", "Generate misalign, page, SO, and atomic cacheability exceptions."],
            ["Stall/restart", "L1247-L1336", "Prioritize atomic, page-cross, DCACHE, and MMU replay conditions."],
            ["Downstream outputs", "L1342-L1456", "Drive DC valid, ahead signals, monitor, LSIQ/IDU, and HPCP events."],
        ],
        [3.4, 3.5, 9.5],
    )

    doc.add_heading("3. Clock, Reset, and Pipeline Timing", level=1)
    doc.add_heading("3.1 Clock Domains", level=2)
    add_table(
        doc,
        ["Clock", "Registers", "Enable/Source", "Verification Requirement"],
        [
            ["ctrl_ld_clk", "ld_ag_inst_vld", "External control clock", "Flush/reset priority and stall retention shall be checked on every active edge."],
            ["ld_ag_clk", "Instruction attributes, base, offset, replay state", "gated_clk_cell from forever_cpuclk", "An accepted RF instruction shall cause a data clock edge whenever its valid is captured."],
        ],
        [3.0, 4.2, 4.0, 5.2],
    )
    doc.add_paragraph(
        "Open assumption A-CLK-1: ctrl_ld_clk and forever_cpuclk/ld_ag_clk are phase-compatible and use coherent enables. The RTL alone does not prove this relationship."
    )
    doc.add_heading("3.2 Reset and Flush Priority", level=2)
    add_bullet(doc, "cpurst_b is active low and asynchronously clears ld_ag_inst_vld and most instruction-attribute registers.")
    add_bullet(doc, "rtu_yy_xx_flush synchronously clears ld_ag_inst_vld with higher priority than stall or new issue.")
    add_bullet(doc, "ld_ag_base and ld_ag_offset have no explicit reset. Their values are undefined until a valid RF load is captured; consumers shall qualify dependent outputs with valid.")
    doc.add_heading("3.3 Nominal Transaction", level=2)
    add_table(
        doc,
        ["Cycle relation", "Condition", "Expected state/action"],
        [
            ["N", "idu_lsu_rf_pipe3_sel/gateclk_sel = 1 and no AG stall", "RF fields are sampled into AG; AG valid becomes 1."],
            ["N+1 combinational", "AG valid = 1", "VA, masks, MMU request, speculative DCACHE request, exception, and stall logic evaluate."],
            ["N+1", "No stall/restart", "ld_ag_dc_inst_vld = 1 and the instruction advances to DC."],
            ["N+1", "stall/restart", "AG valid/data are retained or replay base/offset are updated according to arbitration."],
        ],
        [3.0, 6.0, 7.4],
    )

    doc.add_heading("4. Interface Specification", level=1)
    doc.add_paragraph(
        "Widths are taken directly from the module declaration. Unless a signal description states otherwise, outputs are combinational functions of the current AG registers and shall be sampled only with the associated valid."
    )
    groups = [
        ("Clock, Reset, and CP0", lambda n: n.startswith("cp0_") or n in {"cpurst_b", "ctrl_ld_clk", "forever_cpuclk", "pad_yy_icg_scan_en"}),
        ("IDU/RF Inputs", lambda n: n.startswith("idu_lsu_rf_")),
        ("MMU Interface", lambda n: n.startswith("mmu_lsu_") or n.startswith("lsu_mmu_")),
        ("DCACHE Interface", lambda n: n.startswith("dcache_arb_") or n.startswith("ag_dcache_arb_")),
        ("RTU and Store-AG Inputs", lambda n: n.startswith("rtu_") or n == "st_ag_iid"),
        ("Load-AG and Downstream Outputs", lambda n: n.startswith("ld_ag_") or n.startswith("lsu_idu_") or n.startswith("lsu_hpcp_")),
    ]
    used: set[str] = set()
    for title, predicate in groups:
        selected = [p for p in ports if predicate(p["name"]) and p["name"] not in used]
        if not selected:
            continue
        used.update(p["name"] for p in selected)
        doc.add_heading(f"4.{len([x for x in used]) and groups.index((title, predicate)) + 1} {title}", level=2)
        add_table(
            doc,
            ["Signal", "Dir", "Width", "RTL line", "Description / validity"],
            [[p["name"], p["direction"], p["width"], p["line"], p["description"]] for p in selected],
            [4.2, 1.2, 1.5, 1.5, 8.0],
        )
    remaining = [p for p in ports if p["name"] not in used]
    if remaining:
        doc.add_heading("4.7 Other Interface Signals", level=2)
        add_table(
            doc,
            ["Signal", "Dir", "Width", "RTL line", "Description / validity"],
            [[p["name"], p["direction"], p["width"], p["line"], p["description"]] for p in remaining],
            [4.2, 1.2, 1.5, 1.5, 8.0],
        )

    doc.add_heading("5. Functional Requirements", level=1)
    requirements = [
        ("FR-01", "AG valid capture", "L582-L591", "Reset or flush clears valid. Otherwise stall or RF select sets valid; no stall and no RF select clears valid.", "Check all priority combinations, including flush with stall and flush with new issue."),
        ("FR-02", "AG data capture", "L626-L692", "When no AG stall and RF data is valid, all listed instruction attributes are captured. During stall, they retain their previous values.", "Randomize every field and compare cycle-accurate hold/load behavior."),
        ("FR-03", "Immediate/register offset", "L728-L765", "Normal load uses sign-extended 12-bit immediate; LDR uses src1 with optional high zero extension; split/replay may replace the offset.", "Cross instruction type, sign, zero-extend, and replay."),
        ("FR-04", "Offset shift", "L713-L720, L789-L792", "shift is interpreted as one-hot select of offset << 0,1,2,3.", "Verify four legal values and inject all illegal encodings."),
        ("FR-05", "VA selection", "L794-L808", "va_ori = base + shifted_offset. A first, boundary, ordinary load that is not LDR selects va_plus; otherwise va_ori is selected.", "Reference-model 64-bit modular arithmetic and split behavior."),
        ("FR-06", "Instruction classes", "L816-L835", "Ordinary load, LR, LDAMO, prefetch, scalar/vector-FLS, and tied-off feature classifications follow the explicit equations.", "Enumerate atomic x inst_type x split x lsfifo x secd."),
        ("FR-07", "Access size", "L842-L868", "BYTE/HALF/WORD/DWORD map to size-minus-one 0/1/3/7 and DC size codes 0/1/2/3.", "Exhaustive four-code check."),
        ("FR-08", "Natural alignment", "L872-L900", "BYTE is always aligned; HALF requires A0=0; WORD A1:0=0; DWORD A2:0=0. SO uses the same alignment result.", "Sweep address low bits for every size."),
        ("FR-09", "16-byte boundary", "L903-L909", "Boundary is detected by adding start offset and size-minus-one. Output boundary is restricted to ordinary load and remains asserted on secd.", "Sweep all 16 starts and four sizes; verify first/second beat."),
        ("FR-10", "Byte masks", "L915-L985", "Masks identify bytes from start to end within a 16-byte window; secd selects the high-side mask.", "Compare against a software mask model for all starts/sizes/secd."),
        ("FR-11", "MMU request/abort", "L1035-L1044", "MMU request valid follows AG valid. Address is base. Abort asserts for page replay, atomic wait, cache stall, no-page misalign, or flush.", "Check each abort cause alone and in priority combinations."),
        ("FR-12", "PA and attributes", "L1047-L1068", "PA combines MMU PPN with selected VA low 12 bits. CA/SO are valid-gated; other attributes pass through.", "Use random PPN/VA and invalid MMU cycles."),
        ("FR-13", "Boundary acceleration", "L1071-L1083", "Acceleration requires boundary, same 4 KiB page, feature enable, first beat, dcache enable, and no prior replay; DC enable additionally requires PA valid/cacheable.", "Cross every qualifier and check load classification."),
        ("FR-14", "Commit match", "L1109-L1119", "Current iid is considered committed when any of three valid commit ports carries the same iid.", "Match each port and simultaneous/nonmatching combinations."),
        ("FR-15", "DCACHE array request", "L1124-L1190", "Tag/data array requests are speculative when AG valid and dcache enabled. Bank enables cover 32-bit banks touched by the access; acceleration selects all banks.", "Sweep start/end banks, boundary/secd, dcache enable, and MMU invalid."),
        ("FR-16", "Exceptions", "L1204-L1241", "Generate no-page misalign, SO misalign, page fault, LDAMO non-CA, and vector-SO access fault per RTL equations. Aggregate expt_vld excludes LDAMO non-CA.", "Directed test each exception plus simultaneous conditions and priority at later stages."),
        ("FR-17", "Atomic restart", "L1247-L1249", "Valid atomic instruction restarts until a matching RTU commit is observed.", "Delay commit, use wrong iid, then matching iid; verify hold and release."),
        ("FR-18", "4 KiB crossing", "L1256-L1282", "Crossing is low-12 carry or illegal high offset extension. Replay is suppressed by no-page misalign.", "Boundary-focused random and formal arithmetic equivalence."),
        ("FR-19", "Stall priority", "L1287-L1316", "Priority is atomic wait, cross-page/LDR replay, DCACHE stall, then MMU stall. Older RF instruction may mask non-atomic stall.", "Pairwise and all-cause tests with iid age wrap cases."),
        ("FR-20", "DC advance", "L1342-L1350", "DC valid requires AG valid and no restart/stall. Borrow address overrides normal PA when borrow valid.", "Check every stall cause and borrow mux."),
        ("FR-21", "Ahead signaling", "L1357-L1382, L1423-L1438", "Ahead prediction is constant 1. Scalar/vector valid is qualified by class, boundary, forwarding-disable, and tied-off vector features.", "Prove cancel/replay for cache/TLB/exception misses."),
        ("FR-22", "Performance events", "L1447-L1456", "Cross-4K and other-stall counters are qualified by valid and exclude uTLB miss/already-DA as defined.", "Generate isolated and overlapping stall classes and count pulses."),
    ]
    add_table(doc, ["ID", "Requirement", "Source", "Normative behavior", "Verification"], [list(x) for x in requirements], [1.3, 2.7, 2.4, 6.0, 4.0])

    doc.add_heading("6. Configuration Matrix", level=1)
    configs = [
        ["cp0_lsu_dcache_en", "0/1", "Enables speculative DCACHE tag/data requests and boundary acceleration precondition.", "Both values; toggle only between transactions unless design confirms dynamic behavior."],
        ["cp0_lsu_cb_aclr_dis", "0/1", "Disables 16-byte boundary acceleration when 1.", "Cross with boundary, PA valid, cacheable, and secd."],
        ["cp0_lsu_da_fwd_dis", "0/1", "Suppresses scalar/vector ahead outputs when 1.", "Toggle for ordinary and vector-FLS load."],
        ["cp0_lsu_mm", "0/1", "Controls whether ordinary-load misalignment is reported without waiting for page attributes.", "Cross with alignment, atomic, SO, and PA valid."],
        ["cp0_yy_clk_en", "0/1", "Global clock enable for the AG data clock.", "Clock-gating and wake-up test."],
        ["cp0_lsu_icg_en", "0/1", "Module-level ICG enable.", "Clock-gating and scan override test."],
        ["pad_yy_icg_scan_en", "0/1", "Scan-mode ICG override.", "DFT-mode clock propagation test."],
    ]
    add_table(doc, ["Configuration", "Values", "Effect", "Required test"], configs, [4.0, 1.6, 6.2, 5.0])

    doc.add_heading("7. Verification Strategy", level=1)
    doc.add_heading("7.1 Reference Model", level=2)
    doc.add_paragraph(
        "The scoreboard shall model 64-bit modular address arithmetic, size-minus-one, natural alignment, 16-byte masks, 4 KiB crossing, and PA concatenation. Protocol state shall include AG valid/data hold, already-crossed replay state, atomic commit wait, and older-RF stall masking."
    )
    doc.add_heading("7.2 Directed Test Set", level=2)
    tests = [
        ["T01", "Reset/idle X", "Release reset with no issue; verify no architectural request/event is accepted despite undefined data fields."],
        ["T02", "Issue and advance", "Single ordinary load, MMU hit, DCACHE selected; verify one-cycle AG residency and DC valid."],
        ["T03", "Flush priority", "Assert flush with new issue and with each stall cause; valid shall clear."],
        ["T04", "Data hold", "Random data fields while stalled; all AG registers shall remain stable except defined replay fields."],
        ["T05", "Address arithmetic", "Positive/negative immediate, LDR src1, zero extension, four shift values, wraparound."],
        ["T06", "Alignment sweep", "All size codes x VA[3:0]."],
        ["T07", "Mask sweep", "All size codes x VA[3:0] x secd; compare 16-bit masks and rot_sel."],
        ["T08", "Bank sweep", "All start/end 4-byte banks, boundary and acceleration states."],
        ["T09", "4 KiB replay", "Addresses around page end with positive/negative offsets and split/LDR cases."],
        ["T10", "Atomic commit", "No commit, wrong iid, each commit port match, simultaneous matches, flush."],
        ["T11", "MMU miss/fault", "pa_vld 0, page fault, CA/SO/BUF/SEC/SH attributes, abort causes."],
        ["T12", "Non-cacheable AMO", "Verify dedicated exception and zero architectural side effects."],
        ["T13", "Boundary acceleration", "Enable/disable, cacheable/non-cacheable, PA valid/invalid, SO, 4 KiB cross."],
        ["T14", "DCACHE backpressure", "Toggle arb select; verify hold, abort, and no duplicate DC transaction."],
        ["T15", "Older RF mask", "Exercise iid age comparator including wrap boundary; verify stall entry selection."],
        ["T16", "Ahead cancellation", "Cache miss, uTLB miss, page fault, boundary, flush, forwarding disable."],
        ["T17", "Illegal shift", "Inject zero-hot and multi-hot encodings; expect assertion or documented safe behavior."],
        ["T18", "vreg bit 6", "Drive bit 6 both ways and validate the intended six-bit contract."],
        ["T19", "Idle MMU stall", "Force stall0 during AG idle; verify contract/assertion prevents phantom valid."],
        ["T20", "Clock gating", "Local/global/module/scan enable transitions with issue, stall, and reset."],
    ]
    add_table(doc, ["Test", "Name", "Stimulus and checks"], tests, [1.4, 3.3, 12.0])
    doc.add_heading("7.3 Constrained-Random", level=2)
    for text in [
        "Random instruction attributes with legal one-hot shift and valid size/type constraints.",
        "Random MMU latency/result and DCACHE arbitration, while respecting or deliberately violating documented assumptions in negative tests.",
        "Random RTU commit timing/iid for atomic requests.",
        "Random flush at every AG state and replay phase.",
        "Scoreboard every accepted RF transaction through advance, restart, exception, or flush termination.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("8. Suggested Assertions", level=1)
    assertions = [
        ["A01", "Flush clears valid", "rtu_yy_xx_flush |=> !ld_ag_inst_vld"],
        ["A02", "Stall holds instruction identity", "ld_ag_stall_vld && !rtu_yy_xx_flush |=> $stable(ld_ag_iid)"],
        ["A03", "DC valid implies AG valid", "ld_ag_dc_inst_vld |-> ld_ag_inst_vld"],
        ["A04", "DC valid excludes restart", "ld_ag_dc_inst_vld |-> !ld_ag_stall_restart"],
        ["A05", "MMU request valid equality", "lsu_mmu_va0_vld == ld_ag_inst_vld"],
        ["A06", "Boundary only ordinary load", "ld_ag_boundary |-> ld_ag_ld_inst"],
        ["A07", "Acceleration legality", "ld_ag_dc_acclr_en |-> mmu_lsu_pa0_vld && ld_ag_page_ca && ld_ag_boundary"],
        ["A08", "Legal shift", "RF accept |-> $onehot(idu_lsu_rf_pipe3_shift)"],
        ["A09", "No idle MMU stall", "!ld_ag_inst_vld && !idu_lsu_rf_pipe3_sel |-> !mmu_lsu_stall0"],
        ["A10", "vreg high bit contract", "RF vector accept |-> !idu_lsu_rf_pipe3_vreg[6]"],
        ["A11", "Array request qualification", "ag_dcache_arb_ld_tag_req |-> ld_ag_inst_vld && cp0_lsu_dcache_en"],
        ["A12", "Mask nonzero for valid size", "ld_ag_inst_vld |-> ld_ag_dc_bytes_vld != 16'h0000"],
    ]
    add_table(doc, ["ID", "Intent", "Property sketch"], assertions, [1.4, 5.0, 10.3])
    doc.add_paragraph("Property sketches are not drop-in SVA. Bind paths, sampling clocks, reset conditions, and helper definitions shall be adapted to the verification environment.")

    doc.add_heading("9. Coverage Model", level=1)
    coverage = [
        ["Instruction", "atomic, inst_type, inst_size, LDR, split, secd, FLS, sign-extend"],
        ["Address", "VA[3:0], page offset near 0/FFF, offset sign, shift one-hot, boundary, cross-4K"],
        ["MMU", "pa_vld, page_fault, CA, SO, BUF, SEC, SH, stall"],
        ["Control", "new issue, stall reason, older-RF mask, flush, commit port/hit"],
        ["Configuration", "dcache_en, aclr_dis, da_fwd_dis, mm, ICG enables"],
        ["Outputs", "mask pattern, bank enable, tag/data index, exception class, DC valid, ahead valid, HPCP event"],
    ]
    add_table(doc, ["Covergroup", "Coverpoints"], coverage, [3.3, 13.4])
    doc.add_heading("9.1 Mandatory Crosses", level=2)
    for text in [
        "inst_size x VA[3:0] x boundary x secd",
        "boundary x pa_vld x page_ca x page_so x cp0_lsu_cb_aclr_dis",
        "cross_4K x offset_sign x shift x LDR/split",
        "exception_class x DCACHE request x DC valid",
        "stall_reason x atomic x commit_hit x older_RF_mask",
        "ahead_valid x cache/TLB outcome x flush/replay",
    ]:
        add_bullet(doc, text)

    doc.add_heading("10. Open Questions and Assumptions", level=1)
    questions = [
        ["Q1", "Is mmu_lsu_stall0 guaranteed to remain 0, or at least 0 when AG is idle?"],
        ["Q2", "Is idu_lsu_rf_pipe3_shift strictly one-hot? What is required for illegal encodings?"],
        ["Q3", "Is idu_lsu_rf_pipe3_vreg[6] reserved and guaranteed 0?"],
        ["Q4", "Is ld_ag_ahead_predict = 1 the final policy, and which downstream mechanism cancels every misprediction?"],
        ["Q5", "May DCACHE tag/data arrays be read before PA valid, cacheability, and exception resolution? Which consumer prevents side effects?"],
        ["Q6", "Is exclusion of ld_ag_expt_ldamo_not_ca from ld_ag_expt_vld intentional for every aggregate-exception consumer?"],
        ["Q7", "Why does boundary load-valid use ld_ag_acclr_en rather than the PA-qualified ld_ag_dc_acclr_en?"],
        ["Q8", "Does every page-fault response assert mmu_lsu_pa0_vld in the same cycle?"],
        ["Q9", "What exact phase and stop-clock relationship is guaranteed between ctrl_ld_clk and ld_ag_clk?"],
        ["Q10", "Are vector mask, FOF, VMB merge, and vector-ahead functions permanently removed or planned for restoration?"],
    ]
    add_table(doc, ["ID", "Question / required contract"], questions, [1.4, 15.3])

    doc.add_heading("Appendix A. Traceability", level=1)
    doc.add_paragraph(
        "The companion ct_lsu_ld_ag.md contains one annotation row for every physical source line. The companion ct_lsu_ld_ag_risk_review.md ranks protocol and implementation risks. The Excel feature list maps verification features to RTL line ranges and test methods."
    )
    doc.add_heading("Appendix B. Known RTL Version Characteristics", level=1)
    for text in [
        "Vector mask block is absent/tied off in this version.",
        "ld_ag_inst_vls, ld_ag_inst_fof, ld_ag_vmb_merge_vld, and ld_ag_dc_vload_ahead_inst_vld are constant 0.",
        "ld_ag_ahead_predict is constant 1.",
        "MMU port0 stall is constant 0 in the matching upstream OpenC910 DUTLB implementation, but not constrained inside this module.",
        "The supplied RTL is byte-identical to the current upstream OpenC910 file at the URL on the cover page as checked on 2026-07-11.",
    ]:
        add_bullet(doc, text)

    # Keep headings with the following paragraph where possible.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
    doc.save(SPEC)


def build_model(lines: list[str], ports: list[dict[str, str]], digest: str) -> None:
    MODEL_DIR.mkdir(parents=True, exist_ok=True)
    model = {
        "module": "ct_lsu_ld_ag",
        "source": SOURCE.name,
        "sha256": digest,
        "physical_lines": len(lines),
        "nonblank_lines": sum(bool(x.strip()) for x in lines),
        "upstream_url": UPSTREAM_URL,
        "sections": [{"start": a, "end": b, "name": c} for a, b, c in SECTIONS],
        "ports": ports,
        "clocking": {
            "ctrl_ld_clk": "ld_ag_inst_vld",
            "ld_ag_clk": "instruction data, base, offset, replay state",
            "reset": "cpurst_b active low; base/offset are intentionally not reset",
        },
        "risk_ids": [f"R{i}" for i in range(1, 14)],
        "open_questions": [f"Q{i}" for i in range(1, 11)],
    }
    MODEL.write_text(json.dumps(model, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    raw = SOURCE.read_bytes()
    digest = hashlib.sha256(raw).hexdigest()
    text = raw.decode("utf-8")
    lines = text.splitlines()
    if len(lines) != 1462:
        raise SystemExit(f"Unexpected source line count: {len(lines)}")
    ports = parse_ports(lines)
    if len(ports) < 100:
        raise SystemExit(f"Port parser found only {len(ports)} ports")
    build_model(lines, ports, digest)
    build_annotation(lines, digest)
    build_spec(lines, ports, digest)
    print(json.dumps({"annotation": str(ANNOTATION), "spec": str(SPEC), "model": str(MODEL), "ports": len(ports)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
