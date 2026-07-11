#!/usr/bin/env python3
"""Read-only verification for the generated documentation artifacts."""

from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ct_lsu_ld_ag.v"


def require(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def main() -> None:
    required = [
        ROOT / "ct_lsu_ld_ag.md",
        ROOT / "ld_ag_spec.docx",
        ROOT / "ld_ag_spec.doc",
        ROOT / "ld_ag_feature_list.xlsx",
        ROOT / "ct_lsu_ld_ag_risk_review.md",
    ]
    for path in required:
        require(path.is_file() and path.stat().st_size > 1000, f"missing or empty: {path.name}")

    source_bytes = SOURCE.read_bytes()
    source_lines = source_bytes.decode("utf-8").splitlines()
    digest = hashlib.sha256(source_bytes).hexdigest()
    require(len(source_lines) == 1462, "RTL physical line count changed")
    require(digest == "7c62256e615d66ed54bf96e10373e2636206832cc33d5749e789257f9bb8a010", "RTL hash changed")

    annotation = (ROOT / "ct_lsu_ld_ag.md").read_text(encoding="utf-8")
    markers = [int(x) for x in re.findall(r"SRC-LINE:(\d{4})", annotation)]
    require(markers == list(range(1, 1463)), "annotation markers are missing, duplicated, or out of order")
    for heading in ["虚拟地址生成", "MMU 接口", "异常生成", "stall、restart 与优先级", "阅读注意事项"]:
        require(heading in annotation, f"annotation heading missing: {heading}")

    risk = (ROOT / "ct_lsu_ld_ag_risk_review.md").read_text(encoding="utf-8")
    for risk_id in [f"R{i}" for i in range(1, 14)]:
        require(re.search(rf"\b{risk_id}\b", risk) is not None, f"risk missing: {risk_id}")
    for question_id in [f"{i}." for i in range(1, 11)]:
        require(question_id in risk, f"risk-review question missing: {question_id}")
    require("未发现无需接口上下文即可断言为必现体系结构错误的 P0 问题" in risk, "risk conclusion missing")

    docx = Document(ROOT / "ld_ag_spec.docx")
    docx_text = "\n".join(p.text for p in docx.paragraphs)
    docx_text += "\n" + "\n".join(cell.text for table in docx.tables for row in table.rows for cell in row.cells)
    for heading in [
        "1. Purpose and Scope",
        "4. Interface Specification",
        "5. Functional Requirements",
        "7. Verification Strategy",
        "10. Open Questions and Assumptions",
    ]:
        require(heading in docx_text, f"DOCX heading missing: {heading}")

    port_names = []
    pattern = re.compile(r"^(?:input|output)\s+(?:\[[^\]]+\]\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*;")
    for line in source_lines:
        match = pattern.match(line.strip())
        if match:
            port_names.append(match.group(1))
    missing_ports = [name for name in port_names if name not in docx_text]
    require(not missing_ports, f"DOCX interface table missing ports: {missing_ports[:10]}")

    doc_header = (ROOT / "ld_ag_spec.doc").read_bytes()[:8]
    require(doc_header == bytes.fromhex("D0CF11E0A1B11AE1"), "ld_ag_spec.doc is not a genuine OLE Compound Word document")

    print(
        "PASS artifacts: "
        f"rtl_lines={len(source_lines)}, annotation_markers={len(markers)}, "
        f"ports={len(port_names)}, docx_tables={len(docx.tables)}, risks=13, questions=10"
    )


if __name__ == "__main__":
    main()
